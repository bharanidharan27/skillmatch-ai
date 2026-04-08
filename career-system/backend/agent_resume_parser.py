"""
agent_resume_parser.py – Agent 1: Resume Parser.

Uses spaCy (en_core_web_sm) + rule-based extraction to parse structured
information from preprocessed resumes:
  - Skills (via dictionary matching)
  - Years of experience (regex patterns)
  - Education level
  - Job titles mentioned
  - Tools / technologies

Skills are scored by frequency, recency (later sections weighted higher),
and section-based weights.

File ingestion (parse_file / extract_text_from_bytes):
  - PDF  : pdfplumber
  - DOCX : python-docx  (paragraphs + table cells)
  - Image: pytesseract + Pillow (OCR)
  - Text : plain UTF-8 string
"""

from __future__ import annotations

import io
import re
from pathlib import Path
from typing import Dict, List, Optional, Union

import spacy

from config import PARAMS
from skill_dictionary import SkillLookup, load_dictionary

# Supported file extensions
_IMAGE_EXTS = {".jpg", ".jpeg", ".png", ".bmp", ".tiff", ".tif", ".webp"}

# ── Lazy spaCy loading ──────────────────────────────────────────────────────

_NLP: Optional[spacy.language.Language] = None


def _get_nlp() -> spacy.language.Language:
    global _NLP
    if _NLP is None:
        _NLP = spacy.load("en_core_web_sm", disable=["ner", "textcat"])
    return _NLP


# ── Regex patterns ──────────────────────────────────────────────────────────

_YEARS_EXP_RE = re.compile(
    r"(\d{1,2})\s*\+?\s*(?:years?|yrs?)\s+(?:of\s+)?(?:\w+\s+){0,3}(?:experience|exp)",
    re.IGNORECASE,
)

# Match date ranges like "2020-Present", "Jan 2017 - Dec 2020", "2015 – 2019"
# Captures optional month groups plus the year groups
_DATE_RANGE_RE = re.compile(
    r"((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?)?\s*"
    r"(20\d{2}|19\d{2})"
    r"\s*[\-–—to]+\s*"
    r"((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?)?\s*"
    r"(20\d{2}|19\d{2}|[Pp]resent|[Cc]urrent|[Nn]ow)",
    re.IGNORECASE,
)

# Month name to number
_MONTH_MAP = {
    "jan": 1, "feb": 2, "mar": 3, "apr": 4, "may": 5, "jun": 6,
    "jul": 7, "aug": 8, "sep": 9, "oct": 10, "nov": 11, "dec": 12,
}

_EDUCATION_LEVELS = [
    (re.compile(r"\b(?:ph\.?d|doctorate|doctor of philosophy)\b", re.I), "PhD"),
    (re.compile(r"\b(?:master'?s?|m\.?s\.?|m\.?tech|mba|m\.?b\.?a)\b", re.I), "Masters"),
    (re.compile(r"\b(?:bachelor'?s?|b\.?s\.?|b\.?tech|b\.?e\.?|b\.?a\.?|b\.?sc)\b", re.I), "Bachelors"),
    (re.compile(r"\b(?:associate'?s?|a\.?s\.?|a\.?a\.?|diploma)\b", re.I), "Associates"),
    (re.compile(r"\b(?:high\s+school|ged)\b", re.I), "High School"),
]

_JOB_TITLE_RE = re.compile(
    r"\b((?:senior|sr\.?|junior|jr\.?|lead|principal|staff|chief|head|director|manager|vp|"
    r"associate|intern|consultant)\s+)?"
    r"(software\s+(?:engineer(?:ing)?|developer|architect)|web\s+(?:developer|development\s+intern)|"
    r"java\s+(?:developer|architect)|"
    r"full[\s-]?stack\s+(?:developer|engineer)|front[\s-]?end\s+(?:developer|engineer)|"
    r"back[\s-]?end\s+(?:developer|engineer)|"
    r"data\s+(?:engineer|scientist|analyst)|devops\s+engineer|"
    r"network\s+(?:engineer|administrator)|systems?\s+administrator|"
    r"database\s+administrator|dba|business\s+analyst|project\s+manager|"
    r"scrum\s+master|product\s+manager|qa\s+(?:engineer|analyst)|"
    r"test\s+(?:engineer|analyst)|etl\s+developer|bi\s+(?:developer|analyst)|"
    r"recruiter|talent\s+acquisition|ios\s+developer|android\s+developer|"
    r"cloud\s+(?:engineer|architect)|solutions?\s+architect|"
    r"technical\s+(?:lead|architect|writer|recruiter)|ui/ux\s+designer|"
    r"security\s+(?:engineer|analyst)|machine\s+learning\s+engineer|"
    r"data\s+warehouse\s+(?:developer|architect)|"
    r"(?:software|web|it|hardware)\s+(?:engineering|development|support)\s+(?:intern|co-?op)|"
    r"website\s+(?:manager|administrator|developer)|"
    r"(?:it|help\s*desk)\s+(?:support|technician|specialist|administrator))\b",
    re.IGNORECASE,
)


# ── Core parser ─────────────────────────────────────────────────────────────


class ResumeParser:
    """Agent 1: Parses structured information from a preprocessed resume.

    Supports three entry points:
      parse(preprocessed)          – from a dict produced by preprocessing.py
      parse_file(path_or_bytes, filename)  – from a raw file (PDF/DOCX/image)
      extract_text_from_bytes(contents, filename) – text extraction only
    """

    def __init__(self, skill_lookup: Optional[SkillLookup] = None) -> None:
        self.skill_lookup = skill_lookup or SkillLookup()
        self.section_weights = PARAMS.section_weights

    # ── Public ───────────────────────────────────────────────────────────

    # -- File ingestion entry point ---------------------------------------

    @staticmethod
    def extract_text_from_bytes(contents: bytes, filename: str) -> str:
        """Extract plain text from PDF, DOCX, or image bytes.

        Args:
            contents: Raw file bytes.
            filename: Original filename, used to determine file type.

        Returns:
            Extracted text as a UTF-8 string.

        Raises:
            ValueError: If the file type is not supported.
        """
        ext = Path(filename).suffix.lower()

        if ext == ".docx":
            try:
                import docx  # python-docx
                doc = docx.Document(io.BytesIO(contents))
                # Extract paragraphs (headings, body text, bullet points)
                lines = [p.text for p in doc.paragraphs if p.text.strip()]
                # Also extract table cells — many resume templates use tables
                # for skills grids, contact info, etc.
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            cell_text = cell.text.strip()
                            if cell_text and cell_text not in lines:
                                lines.append(cell_text)
                return "\n".join(lines).strip()
            except Exception as exc:
                raise ValueError(f"Could not read DOCX '{filename}': {exc}") from exc

        elif ext == ".pdf":
            try:
                import pdfplumber
                with pdfplumber.open(io.BytesIO(contents)) as pdf:
                    pages = [page.extract_text() or "" for page in pdf.pages]
                return "\n".join(pages).strip()
            except Exception as exc:
                raise ValueError(f"Could not read PDF '{filename}': {exc}") from exc

        elif ext in _IMAGE_EXTS:
            try:
                import pytesseract
                from PIL import Image
                img = Image.open(io.BytesIO(contents))
                if img.mode not in ("L", "RGB"):
                    img = img.convert("RGB")
                return pytesseract.image_to_string(img).strip()
            except Exception as exc:
                raise ValueError(f"Could not OCR image '{filename}': {exc}") from exc

        else:
            raise ValueError(
                f"Unsupported file type '{ext}'. "
                "Supported formats: PDF, DOCX, JPG, PNG, BMP, TIFF, WebP."
            )

    def parse_file(
        self,
        source: Union[bytes, str, Path],
        filename: Optional[str] = None,
    ) -> Dict:
        """Parse a resume directly from a file path or raw bytes.

        This is the primary entry point for the agent when ingesting uploaded
        files. It runs the full 3-step pipeline internally:
          1. Extract text from PDF / DOCX / image
          2. Preprocess (section segmentation, cleaning)
          3. Parse (skills, experience, education, job titles)

        Args:
            source  : Raw bytes or a file path (str / Path).
            filename: Original filename — required when ``source`` is bytes,
                      used to determine the file type extension.

        Returns:
            Same dict structure as ``parse()``, plus ``extracted_text``.
        """
        from preprocessing import preprocess_single

        # Resolve bytes vs. path
        if isinstance(source, (str, Path)):
            path = Path(source)
            contents = path.read_bytes()
            filename = filename or path.name
        else:
            contents = source
            if not filename:
                raise ValueError("'filename' is required when source is bytes.")

        text = self.extract_text_from_bytes(contents, filename)

        if not text:
            raise ValueError("No text could be extracted from the file.")

        pre = preprocess_single(text)
        result = self.parse(pre)
        result["extracted_text"] = text
        return result

    # -- Standard dict entry point (used by pipeline.py) -----------------

    def parse(self, preprocessed: Dict) -> Dict:
        """Parse a single preprocessed resume dict.

        Args:
            preprocessed: Output of ``preprocessing.preprocess_single``.

        Returns:
            Dict with ``skills``, ``years_experience``, ``education_level``,
            ``job_titles``, ``scored_skills``, and original metadata.
        """
        sections: Dict[str, str] = preprocessed.get("sections", {})
        raw = preprocessed.get("raw_text", "")

        # 1. Skill extraction with section-aware scoring
        scored_skills = self._extract_skills_scored(sections, raw)

        # 2. Years of experience — prefer experience/work sections to avoid
        #    counting education date ranges (e.g. "2017-2021" for a degree).
        exp_text = sections.get("experience", "") or sections.get("work", "") or ""
        years_exp = self._extract_years_experience(exp_text, raw) if exp_text else self._extract_years_experience(raw)

        # 3. Education level
        edu_level = self._extract_education_level(raw)

        # 4. Job titles
        job_titles = self._extract_job_titles(raw)

        return {
            "resume_id": preprocessed.get("resume_id", -1),
            "category": preprocessed.get("category", ""),
            "job_title": preprocessed.get("job_title", ""),
            "skills": sorted(scored_skills.keys()),
            "scored_skills": scored_skills,
            "years_experience": years_exp,
            "education_level": edu_level,
            "job_titles_mentioned": job_titles,
        }

    def parse_batch(self, preprocessed_list: List[Dict]) -> List[Dict]:
        """Parse a batch of preprocessed resumes."""
        return [self.parse(p) for p in preprocessed_list]

    # ── Private helpers ──────────────────────────────────────────────────

    def _extract_skills_scored(
        self, sections: Dict[str, str], raw_text: str
    ) -> Dict[str, float]:
        """Extract skills from each section, applying section weights.

        For non-experience sections:
            score += frequency * section_weight

        For the experience section, recency weighting is applied:
            The section is split into individual job blocks using date-range
            boundaries. Each block receives a recency_multiplier that decays
            linearly with how many years ago the role ended (1.0 for current /
            most-recent, down to recency_min_multiplier for older roles).

            score += frequency * section_weight * recency_multiplier
        """
        skill_scores: Dict[str, float] = {}
        exp_section_weight = self.section_weights.get("experience", 0.8)

        for sec_name, sec_text in sections.items():
            if sec_name in ("experience", "work"):
                # Split into job blocks and score each with recency weighting
                self._score_experience_section(
                    sec_text, exp_section_weight, skill_scores
                )
            else:
                weight = self.section_weights.get(
                    sec_name, self.section_weights.get("other", 0.4)
                )
                hits = self.skill_lookup.extract_from_text(sec_text)
                for skill, count in hits.items():
                    skill_scores[skill] = skill_scores.get(skill, 0.0) + count * weight

        # If no sections were identified, fall back to raw text with default weight
        if not skill_scores:
            hits = self.skill_lookup.extract_from_text(raw_text)
            for skill, count in hits.items():
                skill_scores[skill] = count * 0.6

        return skill_scores

    def _score_experience_section(
        self,
        exp_text: str,
        section_weight: float,
        skill_scores: Dict[str, float],
    ) -> None:
        """Score skills in the experience section with per-job recency weighting.

        Algorithm:
          1. Find all date-range positions in the text (e.g. "Jun 2022 – Aug 2023").
          2. Split the text into job blocks at each date-range boundary.
          3. Compute the end-year of each block to derive its recency multiplier:
               recency = max(min_mult, 1.0 - decay * years_ago)
             where years_ago = current_year - job_end_year.
             Blocks with no detectable date default to the global minimum.
          4. Score skills in each block: freq * section_weight * recency.
        """
        import datetime
        now = datetime.datetime.now()
        decay = PARAMS.recency_decay_per_year
        min_mult = PARAMS.recency_min_multiplier

        # Find all date-range matches and their positions
        date_matches = list(_DATE_RANGE_RE.finditer(exp_text))

        if not date_matches:
            # No dates found — score the whole section at min recency
            hits = self.skill_lookup.extract_from_text(exp_text)
            for skill, count in hits.items():
                skill_scores[skill] = (
                    skill_scores.get(skill, 0.0) + count * section_weight * min_mult
                )
            return

        # Build job blocks: each block starts at a date-range match
        # and ends just before the next one (or the end of the section).
        blocks: List[Dict] = []
        for i, m in enumerate(date_matches):
            start = m.start()
            end = date_matches[i + 1].start() if i + 1 < len(date_matches) else len(exp_text)
            block_text = exp_text[start:end]

            # Parse the end-year / end-month of this role
            _, _, end_month_str, end_year_str = m.groups()
            try:
                if end_year_str.lower() in ("present", "current", "now"):
                    end_year = now.year
                else:
                    end_year = int(end_year_str)
            except (ValueError, AttributeError):
                end_year = None

            blocks.append({"text": block_text, "end_year": end_year})

        # Sort blocks so the most recent is first (highest end_year / None = present)
        blocks.sort(
            key=lambda b: b["end_year"] if b["end_year"] is not None else now.year + 1,
            reverse=True,
        )

        for block in blocks:
            end_year = block["end_year"]
            if end_year is None:
                years_ago = 0  # treat as current
            else:
                years_ago = max(0, now.year - end_year)

            recency = max(min_mult, 1.0 - decay * years_ago)

            hits = self.skill_lookup.extract_from_text(block["text"])
            for skill, count in hits.items():
                skill_scores[skill] = (
                    skill_scores.get(skill, 0.0) + count * section_weight * recency
                )

    @staticmethod
    def _extract_years_experience(text: str, full_text: str = "") -> Optional[int]:
        """Extract years of experience from explicit mentions or date ranges.

        Args:
            text: The text to search (ideally the experience section only).
            full_text: Full resume text used as a broader fallback for explicit mentions.
        """
        # Try explicit mentions first: "5 years of experience" — search full text too
        for t in (text, full_text):
            matches = _YEARS_EXP_RE.findall(t)
            if matches:
                return max(int(m) for m in matches)

        # Fall back to computing from date ranges using month-level precision
        import datetime
        now = datetime.datetime.now()
        total_months = 0
        for start_month_str, start_year_str, end_month_str, end_year_str in _DATE_RANGE_RE.findall(text):
            try:
                start_year = int(start_year_str)
                start_month = 1  # default to January
                if start_month_str:
                    prefix = start_month_str[:3].lower().rstrip('.')
                    start_month = _MONTH_MAP.get(prefix, 1)

                if end_year_str.lower() in ("present", "current", "now"):
                    end_year = now.year
                    end_month = now.month
                else:
                    end_year = int(end_year_str)
                    end_month = 12  # default to December
                    if end_month_str:
                        prefix = end_month_str[:3].lower().rstrip('.')
                        end_month = _MONTH_MAP.get(prefix, 12)

                months = (end_year - start_year) * 12 + (end_month - start_month)
                if 0 < months <= 480:  # sanity: up to 40 years
                    total_months += months
            except ValueError:
                continue

        if total_months > 0:
            # Round up: 3 months -> 1 year, 14 months -> 2 years
            import math
            return max(1, math.ceil(total_months / 12))
        return None

    @staticmethod
    def _extract_education_level(text: str) -> str:
        """Return the highest education level detected."""
        for pat, level in _EDUCATION_LEVELS:
            if pat.search(text):
                return level
        return "Unknown"

    @staticmethod
    def _extract_job_titles(text: str) -> List[str]:
        """Extract unique job titles from the text."""
        titles = set()
        for m in _JOB_TITLE_RE.finditer(text):
            title = m.group(0).strip()
            titles.add(title.title())
        return sorted(titles)


# ── CLI helper ──────────────────────────────────────────────────────────────
if __name__ == "__main__":
    import sys
    from preprocessing import load_dataset, preprocess_single
    from config import COLUMN_TEXT, COLUMN_CATEGORY, COLUMN_JOB_TITLE

    parser = ResumeParser()

    # If a file path is passed as argument, test parse_file()
    if len(sys.argv) > 1:
        file_path = Path(sys.argv[1])
        print(f"Parsing file: {file_path}")
        result = parser.parse_file(file_path)
    else:
        df = load_dataset()
        row = df.iloc[0]
        pre = preprocess_single(row[COLUMN_TEXT], row[COLUMN_CATEGORY], row[COLUMN_JOB_TITLE], 0)
        result = parser.parse(pre)

    print(f"Skills ({len(result['skills'])}): {result['skills'][:15]}")
    print(f"Years exp: {result['years_experience']}")
    print(f"Education: {result['education_level']}")
    print(f"Job titles: {result['job_titles_mentioned'][:5]}")
    print(f"Top scored skills:")
    for sk, sc in sorted(result["scored_skills"].items(), key=lambda x: -x[1])[:10]:
        print(f"  {sk}: {sc:.2f}")
